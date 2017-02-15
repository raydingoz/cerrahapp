from docx import Document
from urllib.request import urlretrieve
import time
import sys
import os
import comtypes.client


def DosyaIndir():                   #download file from internet and save same file
    print("indirme başladı...")
    url = "http://194.27.141.99/egitim_ogretim/ders/2016_2017/16_17_dersprog_ou.doc"
    urlretrieve(url, "16_17_dersprog_ou.doc")
    print("indirme tamamlandı")

def DosyayıCevir():                 #convert .doc file to .docx
    print("çevirme başladı...")
    wdFormatPDF = 16
    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open("D:\\Python\\16_17_dersprog_ou.doc")
    doc.SaveAs("D:\\Python\\16_17_dersprog_ou.docx", FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()
    print("çevirme tamamlandı")

def Yukle():                        #read table from .docx and do something
    document = Document()
    document = Document('16_17_dersprog_ou.docx')
    table = document.tables[1]      #i need second table
    data = []
    keys = None
    prog = 0
    for i, row in enumerate(table.rows):
        sureilk = time.time()
        text = (cell.text for cell in row.cells)
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        print()
        amfi = row_data.__getitem__("Amfi")
        if amfi == "İNG3":
            amfi = "İNG"
        dersno = row_data.__getitem__("Ders No")
        tarih = row_data.__getitem__("Tarih/saat")
        konu = row_data.__getitem__("Konu")
        hoca = row_data.__getitem__("Öğretim Üyesi")
        ders = row_data.__getitem__("Anabilim Dalı")
        print(" >>>\n",
              amfi, "\n",
              dersno, "\n",
              tarih, "\n",
              konu, "\n",
              hoca, "\n",
              ders, "\n",
              ">>>", "\nYüklendi"
              )
        sureson = time.time()
        sureort = sureson - sureilk
        prog += 1
        sure = (3157 - i) * (sureort)
        print("*----------------------------------------*")
        print(i, "/3157  ", round(prog / 31.57, 2), "%   kalan süre:", round(round(sure, 3) / 60), "dk")
        print("*----------------------------------------*")
    print("----------------------------------------------------------------------")
    print("                     Yükleme Tamamlandı                               ")
    print("----------------------------------------------------------------------")
    
DosyaIndir()
DosyayıCevir()
Yukle()