from docx import Document
from urllib.request import urlretrieve
import requests
import time
import webbrowser
import sys
import datetime
import pymysql.cursors
import comtypes.client


def DosyaIndir():  # download file from internet and save same file
    print("indirme başladı...")
    url = "http://194.27.141.99/egitim_ogretim/ders/2016_2017/16_17_dersprog_ou.doc"
    urlretrieve(url, "16_17_dersprog_ou.doc")
    print("indirme tamamlandı")
    print("__________________")
def DosyayıCevir():  # convert .doc file to .docx
    print("çevirme başladı...")
    wdFormatPDF = 16
    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open("D:\\Python\\16_17_dersprog_ou.doc")
    doc.SaveAs("D:\\Python\\16_17_dersprog_ou.docx", FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()
    print("çevirme tamamlandı")
    print("__________________")
def TabloyuSil():
    print("Silme Başladı...")
    connection = pymysql.connect(host=''******',
                                 user='******',
                                 password='******',
                                 db='*******',
                                 charset='utf8mb4',
                                 cursorclass=pymysql.cursors.DictCursor)
    try:
        with connection.cursor() as cursor:
            # Create a new record
            sql_sil = "TRUNCATE TABLE `ders_listesi`";
            cursor.execute(sql_sil)
        connection.commit()
    finally:
        connection.close()
        print("Silme tamamlandı")
        print("________________")

def DB_Ekle(_amfi,_dersno,_tarih,_konu,_hoca,_ders,_sureilk,_prog,_i):
    connection = pymysql.connect(host=''******',
                                 user='******',
                                 password='******',
                                 db='*******',
                                 charset='utf8mb4',
                                 cursorclass=pymysql.cursors.DictCursor)
    try:
        with connection.cursor() as cursor:
            sql = "INSERT INTO `ders_listesi` (`Amfi`, `Ders No`, `Tarih/saat`, `Konu`, `Öğretim Üyesi`, `Anabilim Dalı`) VALUES (%s,%s,%s,%s,%s,%s);";
            cursor.execute(sql, (_amfi, _dersno,_tarih,_konu,_hoca,_ders))
        connection.commit()
    finally:
        connection.close()
        sureson = time.time()
        sureort = sureson - _sureilk
        sure = (3157 - _i) * (sureort)
        print("*----------------------------------------*")
        print(_i, "/3157  ", round(_i / 31.57, 2), "%  süre:", round(round(sure, 3) / 60), "dk")

def Yukle():  # read table from .docx and do something
    print("yükleme başladı...")
    document = Document()
    document = Document('16_17_dersprog_ou.docx')
    table = document.tables[1]  # i need second table
    data = []
    keys = None
    prog = 0
    for i, row in enumerate(table.rows):
        sureilk = time.time()
        text = (cell.text for cell in row.cells)
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        amfi = row_data.__getitem__("Amfi")
        if amfi == "İNG3":
            amfi = "İNG"
        elif amfi=="İNG2":
            amfi = "MT"
        elif amfi =="İNG1":
            amfi ="RG"
        dersno = row_data.__getitem__("Ders No")
        tarih = row_data.__getitem__("Tarih/saat")      #tarih formatı 2017-03-27 14:30:00 şeklinde olmalı!!!!!!
        tarihX=(tarih.split)()[0]
        saat=(tarih.split)()[1]
        d = datetime.datetime.strptime(tarihX, '%d.%m.%Y')
        tarih=datetime.date.strftime(d, "%y-%m-%d")+" "+saat
        konu = row_data.__getitem__("Konu")
        hoca = row_data.__getitem__("Öğretim Üyesi")
        ders = row_data.__getitem__("Anabilim Dalı")
        DB_Ekle(amfi,dersno,tarih,konu,hoca,ders,sureilk,prog,i)
    print("----------------------------------------------------------------------")
    print("                     Yükleme Tamamlandı                               ")
    print("----------------------------------------------------------------------")

DosyaIndir()
DosyayıCevir()
TabloyuSil()
Yukle()
