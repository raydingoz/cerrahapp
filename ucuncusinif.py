from docx import Document
from urllib.request import urlretrieve
import glob, os
import time
import locale
import datetime
import pymysql.cursors
import comtypes.client
import winsound

##pratiklerin word dosyasındaki tablo sırası
pt_1 = 14
pt_2 = 21
pt_3 = 26
pt_4 = 36
pt_5 = 41


##ders listelerinin word dosyasındaki tablo sırası
dt_1 = 18
dt_2 = 23
dt_3 = 33
dt_4 = 38
dt_5 = 43

##Ders programlarının en güncel halini indirir
def DosyaIndir():  # download file from internet and save same file
    print(" ")
    print(" ________________________________________")
    print(" / 3. sınıf - doc indirme başladı...     \\")
    url = "http://www.ctf.edu.tr/egitim_ogretim/ders/2019_2020/19_20_dersprog3.doc"
    urlretrieve(url, "19_20_dersprog3.doc")
    print(" | indirme tamamlandı                    |")
    
## .doc dosyasını .docx dosyasına çevirmek gerekiyor
def DosyayıCevir():  # convert .doc file to .docx
    print(" |- - - - - - - - - - - - - - - - - - - -|")
    print(" | çevirme başladı...                    |")
    wdFormatPDF = 16
    TEST_FILENAME = os.path.join(os.path.dirname(__file__), '19_20_dersprog3.doc')
    word = comtypes.client.CreateObject('Word.Application')
    for file in glob.glob(TEST_FILENAME):
        doc = word.Documents.Open(file)

    TEST_FILENAME2 = os.path.join(os.path.dirname(__file__), '19_20_dersprog3.docx')
    doc.SaveAs(TEST_FILENAME2, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()
    print(" | 3. sınıf - çevirme tamamlandı         |")
    print(" \_______________________________________/")
    
## Güncellemeden önce eski veritabanını temizlemek gerekli
def TabloyuSil():
    print("Silme Başladı...")
    connection = pymysql.connect(host='cerrahapp.cf',
                                 user='****',
                                 password='****',
                                 db='****',
                                 charset='utf8mb4',
                                 cursorclass=pymysql.cursors.DictCursor)
    try:
        with connection.cursor() as cursor:
            # Create a new record
            sql_sil = "TRUNCATE TABLE `ders_listesi`";
            cursor.execute(sql_sil)
        connection.commit()
    finally:
        connection.close()
        print("Silme tamamlandı")
        print("________________")
        
##eski pratik tablosunu silmek için
def TabloyuSilpratik():
    print("Silme Başladı...")
    connection = pymysql.connect(host='cerrahapp.cf',
                                 user='*',
                                 password='****',
                                 db='****',
                                 charset='utf8mb4',
                                 cursorclass=pymysql.cursors.DictCursor)
    try:
        with connection.cursor() as cursor:
            # Create a new record
            sql_sil = "TRUNCATE TABLE `uygulama_ucuncu_sinif`";
            cursor.execute(sql_sil)
        connection.commit()
    finally:
        connection.close()
        print("Silme tamamlandı")
        print("________________")
        
##Veritabanına ders programını ekleme fonksiyonu, tek tek ekliyor
def DB_Ekle_ders(_amfi,_dersno,_tarih,_konu,_hoca,_ders,_sureilk,_prog,_i):
    connection = pymysql.connect(host='cerrahapp.cf',
                                 user='****',
                                 password='****',
                                 db='****',
                                 charset='utf8mb4',
                                 cursorclass=pymysql.cursors.DictCursor)
    try:
        with connection.cursor() as cursor:
            # Create a new record
            sql = "INSERT INTO `ders_listesi` (`Amfi`, `Ders No`, `Tarih/saat`, `Konu`, `Öğretim Üyesi`, `Anabilim Dalı`) VALUES (%s,%s,%s,%s,%s,%s);";
            cursor.execute(sql, (_amfi, _dersno, _tarih, _konu.replace("'", " "), _hoca.replace("'", " "), _ders.replace("'", " ")))
        connection.commit()
    finally:
        connection.close()
        
##veritabanına pratik ekleme fonksiyonu, tek tek
def DB_Ekle_pratik(_tarih,_sayi,_harf,_deger,_sureilk,_prog,_i):
    connection = pymysql.connect(host='cerrahapp.cf',
                                 user='****',
                                 password='****',
                                 db='****',
                                 charset='utf8mb4',
                                 cursorclass=pymysql.cursors.DictCursor)
    try:
        with connection.cursor() as cursor:
            # Create a new record
            sql = "INSERT INTO `uygulama_ucuncu_sinif` (`tarih`, `sayi`, `harf`, `deger`) VALUES (%s,%s,%s,%s);";
            cursor.execute(sql, (_tarih, _sayi,_harf,_deger))
        connection.commit()
    finally:
        connection.close()
        
##Ders listesini word dosyasındaki tablolardan çekme fonksiyonu, tüm büyü burada
def Yukle(_i):  #EKU
    yenigun=False
    document = Document('19_20_dersprog3.docx')
    table = document.tables[_i]  # i need second table
    keys = None
    data = []
    gun = ""
    dersnoEKU = 0
    dersnoING = 0
    dersnobir = 0
    dersnoiki = 0
    dersnoing = 0
    yariyil = "iki"

    ##her bir dersin kendisine ait kodu var, burada o yapılıyor, 3. kuruun 1. amfi için  1. kurul dersinin nosu, 311 ile başlıyor
    if (_i == dt_1):
        yariyil = "iki"
        dersnobir = 311001
        dersnoiki = 321001
        dersnoing = 331001
    elif (_i == dt_2):
        yariyil = "iki"
        dersnobir = 312001
        dersnoiki = 322001
        dersnoing = 332001
    elif (_i == dt_3):
        yariyil = "iki" #iki
        dersnobir = 313001
        dersnoiki = 323001
        dersnoing = 333001
    elif (_i == dt_4):
        yariyil = "iki"
        dersnobir = 314001
        dersnoiki = 324001
        dersnoing = 334001
    elif (_i == dt_5):
        yariyil = "iki"
        dersnobir = 315001
        dersnoiki = 325001
        dersnoing = 335001
        
    ##her bir satır için işlem yapıyor
    for i, row in enumerate(table.rows):

        sureilk = time.time()
        text = (cell.text for cell in row.cells)
        if i == 0:
            keys = tuple(text)
            continue

        row_data = tuple(text)
        ########################
        if (yariyil=="bir"):
            saat = row_data[0]
            ekuders = row_data[1]  # Periferik ve Merkezi Sinir Sistemi, Tanımlamalar (Doç. Dr. Selman Demirci, Anatomi)
            ingders = row_data[2]

            if ("2019" in saat):
                gun = saat.replace(" Pazartesi", "")
                gun = gun.replace(" Salı", "")
                gun = gun.replace(" Çarşamba", "")
                gun = gun.replace(" Perşembe", "")
                gun = gun.replace(" Cuma", "")
            if ("2020" in saat):
                gun = saat.replace(" Pazartesi", "")
                gun = gun.replace(" Salı", "")
                gun = gun.replace(" Çarşamba", "")
                gun = gun.replace(" Perşembe", "")
                gun = gun.replace(" Cuma", "")
            tarihX = gun + " " + saat
            locale.setlocale(locale.LC_ALL, 'turkish')

            try:
                tarihX = datetime.datetime.strptime(tarihX, '%d %B %Y %H:%M')
            except ValueError:
                tarihX = "--"

            data.append([tarihX, ekuders])

            str1 = ekuders.rsplit('(', 1)
            str2 = ingders.rsplit('(', 1)

            hocaEKU = repr(str1[len(str1) - 1])
            hocaEKU = hocaEKU.split(", ")
            hocaING = repr(str2[len(str2) - 1])
            hocaING = hocaING.split(", ")

            konuEKU = str1[0].replace("'", "")
            konuING = str2[0].replace("'", "")

            ogretim_uyesiEKU = hocaEKU[0].replace("'", "")
            ogretim_uyesiING = hocaING[0].replace("'", "")
            try:
                dersEKU = hocaEKU[-1].replace(")'", "")
            except IndexError:
                dersEKU = "---"
            try:
                dersING = hocaING[-1].replace(")'", "")
            except IndexError:
                dersING = "---"

            if (konuEKU != "Uygulamalar" and konuEKU != "Uygulamalar " and konuEKU != "" and konuEKU != " "):
                dersnoEKU = dersnoEKU + 1
                print("")
                print('\033[91m' + "##########################################" + '\033[0m')
                print('\033[93m' + "ders id: " + '\033[0m' + str(dersnoEKU))
                print('\033[93m' + "ders: " + '\033[0m' + str(tarihX))
                print('\033[93m' + "konu: " + '\033[0m' + konuEKU)
                print('\033[93m' + "hoca: " + '\033[0m' + ogretim_uyesiEKU)
                print('\033[93m' + "ders: " + '\033[0m' + dersEKU)
                print('\033[91m' + "##########################################" + '\033[0m')
                print("")
                DB_Ekle_ders("EKU", dersnoEKU, tarihX, konuEKU, ogretim_uyesiEKU, dersEKU, sureilk, 0, i)

            if (konuING != "Practices" and konuING != "Practices " and konuING != "" and konuING != " "):
                dersnoING = dersnoING + 1
                print("")
                print('\033[91m' + "##########################################" + '\033[0m')
                print('\033[93m' + "ders idING: " + '\033[0m' + str(dersnoING))
                print('\033[93m' + "dersING: " + '\033[0m' + str(tarihX))
                print('\033[93m' + "konuING: " + '\033[0m' + konuING)
                print('\033[93m' + "hocaING: " + '\033[0m' + ogretim_uyesiING)
                print('\033[93m' + "dersING: " + '\033[0m' + dersING)
                print('\033[91m' + "##########################################" + '\033[0m')
                print("")
                DB_Ekle_ders("İng3", dersnoING, tarihX, konuING, ogretim_uyesiING, dersING, sureilk, 0, i)
                ######################################################################################
        elif (yariyil=="iki"):
            saat = row_data[0]
            birders = row_data[1]  # Periferik ve Merkezi Sinir Sistemi, Tanımlamalar (Doç. Dr. Selman Demirci, Anatomi) örnek ders satırı böyle
            ikiders = row_data[2]
            ingders = row_data[3]

            if ("2019" in saat):
                gun = saat.replace(" Pazartesi", "")
                gun = gun.replace(" Salı", "")
                gun = gun.replace(" Çarşamba", "")
                gun = gun.replace(" Perşembe", "")
                gun = gun.replace(" Cuma", "")
            if ("2020" in saat):
                gun = saat.replace(" Pazartesi", "")
                gun = gun.replace(" Salı", "")
                gun = gun.replace(" Çarşamba", "")
                gun = gun.replace(" Perşembe", "")
                gun = gun.replace(" Cuma", "")
            tarihX = gun + " " + saat
            locale.setlocale(locale.LC_ALL, 'turkish')
            try:
                tarihX = datetime.datetime.strptime(tarihX, '%d %B %Y %H:%M') ##tarih formati vs
            except ValueError:
                tarihX = "--"

            str1 = birders.rsplit('(', 1)
            str2 = ikiders.rsplit('(', 1)
            str3 = ingders.rsplit('(', 1)

            hocabir = repr(str1[len(str1) - 1])
            hocabir = hocabir.split(", ")

            hocaiki = repr(str2[len(str2) - 1])
            hocaiki = hocaiki.split(", ")

            hocaing = repr(str3[len(str3) - 1])
            hocaing = hocaing.split(", ")

            konubir = str1[0].replace("'", "")
            konuiki = str2[0].replace("'", "")
            konuing = str3[0].replace("'", "")

            ogretim_uyesibir = hocabir[0].replace("'", "")
            ogretim_uyesiiki = hocaiki[0].replace("'", "")
            ogretim_uyesiing = hocaing[0].replace("'", "")
            try:
                dersbir = hocabir[-1].replace(")'", "")
            except IndexError:
                dersbir = "---"
            try:
                dersiki = hocaiki[-1].replace(")'", "")
            except IndexError:
                dersiki = "---"
            try:
                dersing = hocaing[-1].replace(")'", "")
            except IndexError:
                dersing = "---"

            if (konubir != "Uygulamalar" and konubir != "Uygulamalar " and konubir != "" and konubir != " "):
                dersnobir = dersnobir + 1
                print("")
                print('\033[91m' + "##########################################" + '\033[0m')
                print('\033[93m' + "ders id 1: " + '\033[0m' + str(dersnobir))
                print('\033[93m' + "tarih 1: " + '\033[0m' + str(tarihX))
                print('\033[93m' + "konu 1: " + '\033[0m' + konubir)
                print('\033[93m' + "hoca 1: " + '\033[0m' + ogretim_uyesibir)
                print('\033[93m' + "ders 1: " + '\033[0m' + dersbir)
                print('\033[91m' + "##########################################" + '\033[0m')
                print("")
                DB_Ekle_ders("1", dersnobir, tarihX, konubir, ogretim_uyesibir, dersbir, sureilk, 0, i)

            if (konuiki != "Uygulamalar" and konuiki != "Uygulamalar " and konuiki != "" and konuiki != " " ):
                dersnoiki = dersnoiki + 1
                print("")
                print('\033[91m' + "##########################################" + '\033[0m')
                print('\033[93m' + "ders id 2: " + '\033[0m' + str(dersnoiki))
                print('\033[93m' + "tarih 2: " + '\033[0m' + str(tarihX))
                print('\033[93m' + "konu 2: " + '\033[0m' + konuiki)
                print('\033[93m' + "hoca 2: " + '\033[0m' + ogretim_uyesiiki)
                print('\033[93m' + "ders 2: " + '\033[0m' + dersiki)
                print('\033[91m' + "##########################################" + '\033[0m')
                print("")
                DB_Ekle_ders("2", dersnoiki, tarihX, konuiki, ogretim_uyesiiki, dersiki, sureilk, 0, i)
            if (konuing != "Practices" and konuing != "Practices " and konuing != "" and konuing != " "):
                dersnoing = dersnoing + 1
                print("")
                print('\033[91m' + "##########################################" + '\033[0m')
                print('\033[93m' + "ders id ing: " + '\033[0m' + str(dersnoing))
                print('\033[93m' + "tarih ing: " + '\033[0m' + str(tarihX))
                print('\033[93m' + "konu ing: " + '\033[0m' + konuing)
                print('\033[93m' + "hoca ing: " + '\033[0m' + ogretim_uyesiing)
                print('\033[93m' + "ders ing: " + '\033[0m' + dersing)
                print('\033[91m' + "##########################################" + '\033[0m')
                print("")
                DB_Ekle_ders("İng3", dersnoing, tarihX, konuing, ogretim_uyesiing, dersing, sureilk, 0, i)
                
##pratikler word dosyasından çekiliyor
def Pratik(_i):
    document = Document('19_20_dersprog3.docx')
    table = document.tables[_i]  # i need second table
    keys = None
    data = []


    # pratik: 13 , 19 , 23 , 32 , 36

    for i, row in enumerate(table.rows):
        a=2
        sureilk = time.time()
        text = (cell.text for cell in row.cells)
        if i == 0:
            keys = tuple(text)
            continue
        row_data = tuple(text)
        ########################
        #No/Tarih/Anatomi/Fizyoloji/Biyokimya/Histoloji/Patoloji/Klinik
        print("")
        print('\033[91m' + "##########################################" + '\033[0m')
        while a < len(keys):
            pratiktarih=row_data[1]

            pratiktarih = pratiktarih.replace("** ","")
            pratiktarih = pratiktarih.replace("**","")
            pratiktarih = pratiktarih.replace("* ","")
            pratiktarih = pratiktarih.replace("*","")

            locale.setlocale(locale.LC_ALL, 'turkish')
            try:
                pratiktarih = str(datetime.datetime.strptime(pratiktarih, '%d.%m.%Y'))
            except ValueError:
                pratiktarih = "---"

            if("," in row_data[a]):
                b = 0
                #pratik içinde virgül var
                tumgrup=row_data[a].split(",")    #1,2 / 2,4 / 2a,2/ 1b,2a
                grup1=len(tumgrup)
                while b < grup1:
                    sayiB = tumgrup[b]  #2/2a/1b
                    pratik = keys[a]
                    pratikb = pratik.rsplit(' / ')
                    if "sınav" in sayiB:
                        pratikb[0] = pratikb[0] + " (Sınav)"
                        pratikb[1] = pratikb[1] + " (exam)"
                        sayiB = sayiB.replace(" (sınav)", "")
                    if ("1" in sayiB):
                        print("tarih: " + pratiktarih + " " + "sayı: " + sayiB + " pratik: " + pratikb[1])
                        if (("a" in sayiB) or ("b" in sayiB)):
                            DB_Ekle_pratik(pratiktarih, sayiB[0], (sayiB[1]).upper(), pratikb[1], sureilk, 0, i)
                        else:
                            DB_Ekle_pratik(pratiktarih, sayiB, "A", pratikb[1], sureilk, 0, i)
                            DB_Ekle_pratik(pratiktarih, sayiB, "B", pratikb[1], sureilk, 0, i)
                    else:
                        print("tarih: " + pratiktarih + " " + "sayı: " + sayiB + " pratik: " + pratikb[0])
                        if (("a" in sayiB) or ("b" in sayiB)):
                            DB_Ekle_pratik(pratiktarih, sayiB[0], (sayiB[1]).upper(), pratikb[0], sureilk, 0, i)
                        else:
                            DB_Ekle_pratik(pratiktarih, sayiB, "A", pratikb[0], sureilk, 0, i)
                            DB_Ekle_pratik(pratiktarih, sayiB, "B", pratikb[0], sureilk, 0, i)
                    b=b+1
            else:
                # pratik içinde virgül yok tek değer
                sayi = row_data[a]   #1 / 2 / 2a / 1b
                if sayi != "-" and sayi != "":

                    pratik = keys[a]
                    if pratik == "Biyoistatistik/ Biostatistics":
                        pratik = "Biyoistatistik / Biostatistics"
                    pratikb = pratik.rsplit(' / ')
                    if "sınav" in sayi:
                        pratikb[0] = pratikb[0] + " (Sınav)"
                        pratikb[1] = pratikb[1] + " (exam)"
                        sayi = sayi.replace(" (sınav)", "")
                    try:
                        if ("1" in sayi):
                            print("tarih: " + pratiktarih + " " + "sayı: " + sayi + " pratik: " + pratikb[1])
                            if (("a" in sayi) or ("b" in sayi)):
                                DB_Ekle_pratik(pratiktarih, sayi[0], (sayi[1]).upper(), pratikb[1], sureilk, 0, i)
                            else:
                                DB_Ekle_pratik(pratiktarih, sayi, "A", pratikb[1], sureilk, 0, i)
                                DB_Ekle_pratik(pratiktarih, sayi, "B", pratikb[1], sureilk, 0, i)
                        else:
                            print("tarih: " + pratiktarih + " " + "sayı: " + sayi + " pratik: " + pratikb[0])
                            if (("a" in sayi) or ("b" in sayi)):
                                DB_Ekle_pratik(pratiktarih, sayi[0], (sayi[1]).upper(), pratikb[0], sureilk, 0, i)
                            else:
                                DB_Ekle_pratik(pratiktarih, sayi, "A", pratikb[0], sureilk, 0, i)
                                DB_Ekle_pratik(pratiktarih, sayi, "B", pratikb[0], sureilk, 0, i)
                    except IndexError:
                        print("")
            a=a+1
        print('\033[91m' + "##########################################" + '\033[0m')
        print("")

        #DB_Ekle("EKU", dersnoEKU, tarihX, konuEKU, ogretim_uyesiEKU, dersEKU, sureilk, 0, i)
        
##tablo numaralarının eşleşmesi doğru mu kontrol ediliyor
def Tablolar(_i):
    document = Document('19_20_dersprog3.docx')
    table = document.tables[_i]  # i need second table
    keys = None
    data = []


    # pratik: 13 , 19 , 23 , 32 , 36

    for i, row in enumerate(table.rows):
        a=2
        sureilk = time.time()
        text = (cell.text for cell in row.cells)
        if i == 0:
            keys = tuple(text)
            continue
        row_data = tuple(text)
        data.append(row_data)
    print(str(data)[:30])
        #DB_Ekle("EKU", dersnoEKU, tarihX, konuEKU, ogretim_uyesiEKU, dersEKU, sureilk, 0, i)
        
def DerslistesiniSil():
    TabloyuSil()
def ders_guncelle():
            Yukle(dt_1)
            Yukle(dt_2)
            Yukle(dt_3)
            Yukle(dt_4)
            Yukle(dt_5)
def pratik_guncelle():
    TabloyuSilpratik()
    Pratik(pt_1)
    Pratik(pt_2)
    Pratik(pt_3)
    Pratik(pt_4)
    Pratik(pt_5)
def indir():
    DosyaIndir()
    DosyayıCevir()
    

def ses ():
    duration = 100  # millisecond
    freq1 = 440  # Hz
    freq2 = 540  # Hz
    freq3 = 640  # Hz
    winsound.Beep(freq1, duration)
    winsound.Beep(freq2, duration)
    winsound.Beep(freq3, duration)
    winsound.Beep(freq1, duration)
    winsound.Beep(freq1, duration)
    winsound.Beep(freq2, duration)
    winsound.Beep(freq3, duration)
    winsound.Beep(freq1, duration)
def TabloPratiklerYaz():
    print("******2. Sınıf Pratik Programları ********")
    Tablolar(pt_1)
    Tablolar(pt_2)
    Tablolar(pt_3)
    Tablolar(pt_4)
    Tablolar(pt_5)
    print("*********************")
def TabloDerslerYaz():
    print("******3. Sınıf Ders Programları *********")
    Tablolar(dt_1)
    Tablolar(dt_2)
    Tablolar(dt_3)
    Tablolar(dt_4)
    Tablolar(dt_5)
    print("*********************")
